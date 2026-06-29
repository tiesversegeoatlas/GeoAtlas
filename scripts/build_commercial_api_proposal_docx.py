from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "GEOATLAS_COMMERCIAL_API_PROPOSAL.md"
OUTPUT = ROOT / "docs" / "GeoAtlas_Commercial_API_Proposal.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "17212B"
MUTED = "667085"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BORDER = "CBD5E1"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=None, bold=None, italic=None, color=INK, name="Aptos") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_inline_markdown(paragraph, text: str, *, size=10.5, color=INK) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=size - 0.5, color=DARK_BLUE, name="Consolas")
        elif re.fullmatch(r"\[.*?\]\(.*?\)", part):
            match = re.fullmatch(r"\[(.*?)\]\((.*?)\)", part)
            run = paragraph.add_run(match.group(1) if match else part)
            set_font(run, size=size, color=BLUE)
            run.underline = True
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size, color=color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.22

    for name, size, color, before, after in (
        ("Title", 27, NAVY, 0, 10),
        ("Subtitle", 14, MUTED, 0, 18),
        ("Heading 1", 16, BLUE, 18, 9),
        ("Heading 2", 13, BLUE, 13, 6),
        ("Heading 3", 11.5, DARK_BLUE, 9, 4),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1", "Heading 2"} else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.4)

    header = section.first_page_header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("INTERNAL PROPOSAL  |  GEOATLAS")
    set_font(run, size=8.5, bold=True, color=MUTED)

    doc.add_paragraph().paragraph_format.space_after = Pt(45)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(15)
    run = kicker.add_run("COMMERCIALIZATION PROPOSAL")
    set_font(run, size=10, bold=True, color=BLUE)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("GeoAtlas Intelligence API")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "A standalone, AI-enriched geospatial risk-intelligence service"
    )

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [9360])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    add_inline_markdown(
        p,
        "Decision requested: approve a controlled commercial pilot and the proposed INR pricing structure.",
        size=11,
        color=NAVY,
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(40)
    for label, value in (
        ("Prepared by", "GeoAtlas API Development"),
        ("Submitted for", "Management review and approval"),
        ("Date", "June 29, 2026"),
        ("Status", "Internal proposal"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        label_run = p.add_run(f"{label}: ")
        set_font(label_run, size=10, bold=True, color=MUTED)
        value_run = p.add_run(value)
        set_font(value_run, size=10, color=INK)

    footer = section.first_page_footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Confidential - for internal management review")
    set_font(run, size=8.5, color=MUTED)
    doc.add_page_break()


def add_running_furniture(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.clear()
    run = header.add_run("GeoAtlas Commercial Intelligence API")
    set_font(run, size=8.5, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.clear()
    add_page_number(footer)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def table_widths(headers: list[str]) -> list[int]:
    count = len(headers)
    if count == 2:
        return [3000, 6360]
    if count == 3:
        return [2600, 3380, 3380]
    if count == 4:
        if "Monthly price" in headers or "Support" in headers:
            return [1900, 2200, 2560, 2700]
        return [1850, 2500, 2500, 2510]
    if count == 5:
        return [1500, 1500, 1900, 1500, 2960]
    base = 9360 // count
    widths = [base] * count
    widths[-1] += 9360 - sum(widths)
    return widths


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    widths = table_widths(rows[0])
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        for col_index in range(cols):
            cell = table.cell(row_index, col_index)
            cell.text = ""
            if row_index == 0:
                set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            value = values[col_index] if col_index < len(values) else ""
            add_inline_markdown(
                p,
                value,
                size=8.7 if cols >= 4 else 9.2,
                color=WHITE if row_index == 0 else INK,
            )
            for run in p.runs:
                if row_index == 0:
                    run.bold = True
            if col_index > 0 and re.search(r"₹|%|requests/minute|Completed|Required", value):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_markdown_body(doc: Document) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(index for index, line in enumerate(lines) if line == "## Executive summary")
    index = start
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                table = doc.add_table(rows=1, cols=1)
                set_table_geometry(table, [9360])
                cell = table.cell(0, 0)
                set_cell_shading(cell, LIGHT_GRAY)
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run("\n".join(code_lines))
                set_font(run, size=9, color=DARK_BLUE, name="Consolas")
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:], style="Heading 1")
            if stripped == "## Executive summary":
                p.paragraph_format.space_before = Pt(0)
            index += 1
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 2")
            index += 1
            continue
        if stripped.startswith("- "):
            item_lines = [stripped[2:]]
            index += 1
            while index < len(lines):
                nxt = lines[index].strip()
                if (
                    not nxt
                    or nxt.startswith("#")
                    or nxt.startswith("|")
                    or nxt.startswith("- ")
                    or nxt.startswith("```")
                    or re.match(r"^\d+\.\s+", nxt)
                ):
                    break
                item_lines.append(nxt)
                index += 1
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, " ".join(item_lines))
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            item_lines = [numbered.group(2)]
            index += 1
            while index < len(lines):
                nxt = lines[index].strip()
                if (
                    not nxt
                    or nxt.startswith("#")
                    or nxt.startswith("|")
                    or nxt.startswith("- ")
                    or nxt.startswith("```")
                    or re.match(r"^\d+\.\s+", nxt)
                ):
                    break
                item_lines.append(nxt)
                index += 1
            p = doc.add_paragraph(style="List Number")
            add_inline_markdown(p, " ".join(item_lines))
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            nxt = lines[index].strip()
            if (
                not nxt
                or nxt.startswith("#")
                or nxt.startswith("|")
                or nxt.startswith("- ")
                or nxt.startswith("```")
                or re.match(r"^\d+\.\s+", nxt)
            ):
                break
            paragraph_lines.append(nxt)
            index += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline_markdown(p, " ".join(paragraph_lines))


def build() -> None:
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_running_furniture(doc)
    add_markdown_body(doc)

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)

    core = doc.core_properties
    core.title = "Proposal for Commercialization of the GeoAtlas Intelligence API"
    core.subject = "Management proposal for the GeoAtlas commercial API"
    core.author = "GeoAtlas API Development"
    core.keywords = "GeoAtlas, API, commercial proposal, pricing, INR"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
